from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "salesforce-cicd-learning-steps.md"
OUTPUT = ROOT / "docs" / "salesforce-cicd-learning-steps.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
GRAY = RGBColor(90, 90, 90)
LIGHT_FILL = "E8EEF5"
CODE_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_style_fill(style, fill):
    p_pr = style.element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = GRAY
    subtitle.paragraph_format.space_after = Pt(14)

    code = doc.styles.add_style("Code Block CICD", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.1)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.05
    set_paragraph_style_fill(code, CODE_FILL)

    table_text = doc.styles.add_style("Table Text CICD", 1)
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.1

    callout = doc.styles.add_style("Callout CICD", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(8)
    set_paragraph_style_fill(callout, "F4F6F9")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    widths = [Inches(2.0), Inches(4.3)] if len(headers) == 2 else [Inches(2.1)] * len(headers)
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.style = "Table Text CICD"
        run = paragraph.add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.style = "Table Text CICD"
            paragraph.add_run(value)
    doc.add_paragraph()


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Salesforce CI/CD Learning Steps")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.style = "Subtitle"
    subtitle.add_run("Panduan belajar dari setup awal, backup, Pull Request, validate, deploy, report Excel, rollback metadata, sampai data restore manual.")

    add_simple_table(
        doc,
        ["Area", "Ringkasan"],
        [
            ("Dummy Dev", "ShankaraOrg"),
            ("Dummy Prod", "trailhead"),
            ("Repository", "Ogifrn09/salesforce-cicd"),
            ("Deploy", "Manual via GitHub Actions dengan confirm_deploy = DEPLOY"),
            ("Rollback", "Manual via destructiveChanges.xml dengan confirm_rollback = ROLLBACK"),
        ],
    )

    note = doc.add_paragraph()
    note.style = "Callout CICD"
    note.add_run("Catatan: ").bold = True
    note.add_run("Rollback metadata tidak otomatis restore data record. Data restore dijalankan manual per object dari CSV backup.")
    doc.add_section(WD_SECTION.NEW_PAGE)


def clean_inline(text):
    return text.replace("`", "")


def add_code_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = "Code Block CICD"
    run = paragraph.add_run(text if text else " ")
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    return paragraph


def add_markdown_content(doc, markdown):
    in_code = False
    code_lines = []

    def flush_code():
        nonlocal code_lines
        for line in code_lines:
            add_code_paragraph(doc, line)
        code_lines = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        bullet_match = re.match(r"^-\s+(.*)", line)
        numbered_match = re.match(r"^\d+\.\s+(.*)", line)

        if bullet_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(clean_inline(bullet_match.group(1)))
        elif numbered_match:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(clean_inline(numbered_match.group(1)))
        else:
            paragraph = doc.add_paragraph()
            paragraph.add_run(clean_inline(line))


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_markdown_content(doc, markdown)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Salesforce CI/CD Learning Steps"
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = GRAY

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
